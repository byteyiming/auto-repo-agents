"""
Format Converter Agent
Converts documentation between different formats (Markdown, HTML, PDF, DOCX)
"""
from typing import Optional, List
from datetime import datetime
from pathlib import Path
from src.agents.base_agent import BaseAgent
from src.utils.file_manager import FileManager
from src.context.context_manager import ContextManager
from src.context.shared_context import AgentType, DocumentStatus, AgentOutput
from src.rate_limit.queue_manager import RequestQueue


class FormatConverterAgent(BaseAgent):
    """
    Format Converter Agent
    
    Converts documentation between formats:
    - Markdown to HTML
    - Markdown to PDF (via HTML)
    - Markdown to DOCX
    - Preserves formatting and structure
    """
    
    def __init__(
        self,
        provider_name: Optional[str] = None,
        model_name: Optional[str] = None,
        rate_limiter: Optional[RequestQueue] = None,
        file_manager: Optional[FileManager] = None,
        api_key: Optional[str] = None,
        **provider_kwargs
    ):
        """Initialize Format Converter Agent"""
        # FormatConverter doesn't actually use LLM, but inherits for consistency
        super().__init__(
            provider_name=provider_name,
            model_name=model_name,
            rate_limiter=rate_limiter,
            api_key=api_key,
            **provider_kwargs
        )
        
        self.file_manager = file_manager or FileManager(base_dir="docs/formats")
        self.supported_formats = ["html", "pdf", "docx"]
    
    def markdown_to_html(self, markdown_content: str) -> str:
        """
        Convert Markdown to HTML
        
        Args:
            markdown_content: Markdown content to convert
        
        Returns:
            HTML content
        """
        try:
            import markdown
            html_content = markdown.markdown(
                markdown_content,
                extensions=['extra', 'codehilite', 'tables']
            )
            # Wrap in proper HTML structure
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Documentation</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            return full_html
        except ImportError:
            # Fallback: basic conversion without markdown library
            print("⚠️  Warning: markdown library not installed, using basic HTML conversion")
            html_content = markdown_content.replace('\n', '<br>\n')
            return f"<html><body>{html_content}</body></html>"
    
    def html_to_pdf(self, html_content: str, output_path: Optional[str] = None) -> str:
        """
        Convert HTML to PDF
        
        Args:
            html_content: HTML content to convert
            output_path: Optional output file path
        
        Returns:
            Path to generated PDF file
        """
        try:
            # Try using weasyprint first (better quality)
            from weasyprint import HTML
            from io import BytesIO
            
            if not output_path:
                output_path = "documentation.pdf"
            
            # Ensure .pdf extension
            if not output_path.endswith('.pdf'):
                output_path = str(Path(output_path).with_suffix('.pdf'))
            
            html_obj = HTML(string=html_content)
            pdf_path = self.file_manager.base_dir / output_path
            html_obj.write_pdf(pdf_path)
            
            return str(pdf_path.absolute())
        except ImportError:
            try:
                # Fallback to pdfkit if available
                import pdfkit
                
                if not output_path:
                    output_path = "documentation.pdf"
                if not output_path.endswith('.pdf'):
                    output_path = str(Path(output_path).with_suffix('.pdf'))
                
                pdf_path = self.file_manager.base_dir / output_path
                pdfkit.from_string(html_content, str(pdf_path))
                
                return str(pdf_path.absolute())
            except ImportError:
                raise ImportError(
                    "PDF conversion requires 'weasyprint' or 'pdfkit'. "
                    "Install with: pip install weasyprint"
                )
    
    def markdown_to_docx(self, markdown_content: str, output_path: Optional[str] = None) -> str:
        """
        Convert Markdown to DOCX
        
        Args:
            markdown_content: Markdown content to convert
            output_path: Optional output file path
        
        Returns:
            Path to generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches
            import re
            
            if not output_path:
                output_path = "documentation.docx"
            if not output_path.endswith('.docx'):
                output_path = str(Path(output_path).with_suffix('.docx'))
            
            doc = Document()
            
            # Parse markdown and convert to docx
            lines = markdown_content.split('\n')
            for line in lines:
                # Headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Lists
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                # Regular paragraphs
                elif line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
            
            docx_path = self.file_manager.base_dir / output_path
            doc.save(str(docx_path))
            
            return str(docx_path.absolute())
        except ImportError:
            raise ImportError(
                "DOCX conversion requires 'python-docx'. "
                "Install with: pip install python-docx"
            )
    
    def convert(
        self,
        markdown_content: str,
        output_format: str,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Convert Markdown content to specified format
        
        Args:
            markdown_content: Markdown content to convert
            output_format: Target format ('html', 'pdf', 'docx')
            output_filename: Optional output filename
        
        Returns:
            Path to converted file
        """
        if output_format.lower() not in self.supported_formats:
            raise ValueError(
                f"Unsupported format: {output_format}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )
        
        if output_format.lower() == 'html':
            html_content = self.markdown_to_html(markdown_content)
            if not output_filename:
                output_filename = "documentation.html"
            file_path = self.file_manager.write_file(output_filename, html_content)
            return file_path
        
        elif output_format.lower() == 'pdf':
            html_content = self.markdown_to_html(markdown_content)
            return self.html_to_pdf(html_content, output_filename)
        
        elif output_format.lower() == 'docx':
            return self.markdown_to_docx(markdown_content, output_filename)
        
        else:
            raise ValueError(f"Format conversion not implemented: {output_format}")
    
    def convert_all_documents(
        self,
        documents: dict,
        formats: List[str],
        project_id: Optional[str] = None,
        context_manager: Optional[ContextManager] = None
    ) -> dict:
        """
        Convert all documents to multiple formats
        
        Args:
            documents: Dict mapping document names to markdown content
            formats: List of target formats (e.g., ['html', 'pdf'])
            project_id: Optional project ID for context
            context_manager: Optional context manager
        
        Returns:
            Dict mapping document names to converted file paths
        """
        results = {}
        
        print(f"🔄 Converting {len(documents)} documents to formats: {', '.join(formats)}")
        
        for doc_name, markdown_content in documents.items():
            doc_results = {}
            
            for fmt in formats:
                try:
                    base_name = Path(doc_name).stem
                    output_filename = f"{base_name}.{fmt}"
                    
                    file_path = self.convert(
                        markdown_content=markdown_content,
                        output_format=fmt,
                        output_filename=output_filename
                    )
                    
                    doc_results[fmt] = file_path
                    print(f"  ✅ Converted {doc_name} to {fmt}")
                    
                except Exception as e:
                    print(f"  ❌ Error converting {doc_name} to {fmt}: {e}")
                    doc_results[fmt] = None
            
            results[doc_name] = doc_results
        
        # Save to context if available
        if project_id and context_manager:
            output = AgentOutput(
                agent_type=AgentType.FORMAT_CONVERTER,
                document_type="format_conversions",
                content=str(results),  # JSON-like string of results
                file_path="",  # Multiple files, no single path
                status=DocumentStatus.COMPLETE,
                generated_at=datetime.now()
            )
            context_manager.save_agent_output(project_id, output)
            print(f"✅ Format conversions saved to shared context (project: {project_id})")
        
        return results

